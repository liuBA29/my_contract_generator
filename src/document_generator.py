# src/document_generator.py

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words
import os, sys
from tkinter import filedialog
from dotenv import load_dotenv
import os

load_dotenv()

USER_COMPANY = os.getenv("USER_COMPANY", 'ООО "Ваша компания"')
USER_COMPANY_SHORT = os.getenv("USER_COMPANY_SHORT", 'ООО "Ваша компания"')
USER_NAME = os.getenv("USER_NAME", 'И.И. Иванов')
AS = os.getenv("AS", 'в лице директора Иванова Ивана Ивановича, действующего ')
AS = AS + ' ' if AS else ''
BASED_ON = os.getenv("BASED_ON", "Устава")
OUR_REQUIZITS = os.getenv("OUR_REQUIZITS", "").replace("\\n", "\n")


#------------------------------------------------

def generate_docx_act(customer, work_list, contract_number, doc_date, act_date,  total_cost):
    doc = Document()

    # Apply styles and margins
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10)
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.5)

    # Header
    p = doc.add_paragraph(f'Приложение 2 к ДОГОВОРУ № {contract_number}  (от {doc_date} года) ', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph(f'Акт сдачи-приемки работ', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(12)
    run.font.bold = True

    # Date
    table = doc.add_table(rows=1, cols=2)
    cell1 = table.cell(0, 0)
    cell1.text = 'г. Минск'
    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1.paragraphs[0].runs[0].font.size = Pt(11)

    cell2 = table.cell(0, 1)
    cell2.text = act_date
    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell2.paragraphs[0].runs[0].font.size = Pt(11)

    # Main text
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p.add_run(
        f"1. {customer[1]}, в лице {customer[9]} {customer[2]}, действующего на основании {customer[3]}, именуемое в дальнейшем ЗАКАЗЧИК, "
        f"с одной стороны, и {USER_COMPANY}, {AS}"
        f"на основании {BASED_ON}, именуемый в дальнейшем "
        f"ИСПОЛНИТЕЛЬ, с другой стороны, составили настоящий Акт о том, что выполненные Исполнителем работы:"
    ).font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 1
    p = doc.add_paragraph('Работы: ',
                          style='Normal')

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    for idx, work in enumerate(work_list, 1):
        p = doc.add_paragraph(f"1.{idx}. {work}.", style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    p=doc.add_paragraph(f'удовлетворяют условиям Договора № {contract_number} от {doc_date} г.')

    # Section 2
    rubles, kopecks = divmod(total_cost, 1)
    rubles = int(rubles)
    kopecks = round(kopecks * 100)  # Округление до целых копеек
    sum_in_words = num2words(rubles, lang='ru', to='cardinal')

    p = doc.add_paragraph(
        f'2. Договорная цена выполненных работ составляет {total_cost:.2f} '
        f'({sum_in_words.capitalize()} белорусских рублей 00 коп.). Без НДС.',
        style='Normal'
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 3

    p = doc.add_paragraph(
            f'3. Исполнитель выполнил свои обязательства в полном объеме.',
            style='Normal'
        )

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 4

    p = doc.add_paragraph(f'4. Заказчик выполненные работы принял, претензий не имеет.', style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 5
    p = doc.add_paragraph(
        '5. Настоящий Акт составлен в 2(двух) экземплярах, один из которых находится у Исполнителя,'
        ' второй - у Заказчика.',
        style='Normal'
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


    # Party details in a table
    recv = doc.add_paragraph('Реквизиты сторон:', style='Normal')
    recv.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    recv.runs[0].font.bold = True
    recv.paragraph_format.space_before = Pt(0)
    recv.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Pt(200)
    table.columns[1].width = Pt(200)

    # Fill in table
    row_cells = table.rows[0].cells
    # Set text and font size for the first cell (Исполнитель)
    p_ispolnitel = row_cells[0].paragraphs[0]

    # Создаем жирный Run для "ИП Панченко К.А."
    run_ispolnitel_bold = p_ispolnitel.add_run(f'{USER_COMPANY_SHORT}\n')
    run_ispolnitel_bold.font.size = Pt(9)
    run_ispolnitel_bold.bold = True

    # Добавляем остальной текст (обычным шрифтом)
    run_ispolnitel = p_ispolnitel.add_run(
        f"{OUR_REQUIZITS}"
    )
    run_ispolnitel.font.size = Pt(9)

    # Set text and font size for the second cell (Заказчик)
    p_zakazchik = row_cells[1].paragraphs[0]

    # Создаем жирный Run для названия заказчика
    run_zakazchik_bold = p_zakazchik.add_run(f'{customer[10]}\n')
    run_zakazchik_bold.font.size = Pt(9)
    run_zakazchik_bold.bold = True

    # Добавляем остальной текст (обычным шрифтом)
    run_zakazchik = p_zakazchik.add_run(
        f'{customer[5]}\n'
        f'IBAN: {customer[8]}\n'
        f'УНП: {customer[6]}, ОКПО: {customer[7]}\n'

    )
    run_zakazchik.font.size = Pt(9)

    # Set column widths
    table.columns[0].width = Pt(150)  # Example width for the first column
    table.columns[1].width = Pt(350)  # Example width for the second column

    # Signatures
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Set column widths
    table.columns[0].width = Pt(150)  # Example width for the first column
    table.columns[1].width = Pt(450)  # Example width for the second column

    cell1 = table.cell(0, 0)
    cell1.text = f'Исполнитель_________({USER_NAME})'
    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1.paragraphs[0].runs[0].font.size = Pt(9)

    cell2 = table.cell(0, 1)
    cell2.text = f'Заказчик_________({customer[4]})'
    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell2.paragraphs[0].runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    p = doc.add_paragraph()







    # Save document with act number

    file_path = filedialog.asksaveasfilename(
        defaultextension=f".docx",
        filetypes=[("Word files", "*.docx")],
        initialfile=f'act_{contract_number}.docx'
    )

    # file_name = f'../docs_out/contract_{contract_number}.docx'
    if file_path:
        # Сохраняем документ
        doc.save(file_path)

        # Получаем папку, в которую пользователь сохранил файл
        saved_folder = os.path.dirname(file_path)
        print("Папка, выбранная пользователем:", saved_folder)

        return file_path
    else:
        return None







def generate_docx(customer, work_list, payment_term, completion, contract_number, location,
                  doc_date, total_cost):
    doc = Document()

    # Apply styles and margins
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10)
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.5)

    # Header
    p = doc.add_paragraph(f'ДОГОВОР № {contract_number}', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.bold = True

    # Date
    table = doc.add_table(rows=1, cols=2)
    cell1 = table.cell(0, 0)
    cell1.text = 'г. Минск'
    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1.paragraphs[0].runs[0].font.size = Pt(10)

    cell2 = table.cell(0, 1)
    cell2.text = doc_date
    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell2.paragraphs[0].runs[0].font.size = Pt(10)

    # Main text
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.add_run(
         f"1. {customer[1]}, в лице {customer[9]} {customer[2]}, действующего на основании {customer[3]}, именуемое в дальнейшем ЗАКАЗЧИК, "
        f"с одной стороны, и {USER_COMPANY}, {AS}"
        f"на основании {BASED_ON}, именуемый в дальнейшем "
        f"ИСПОЛНИТЕЛЬ, с другой стороны, заключили настоящий договор о нижеследующем:"
    ).font.size = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Section 1
    p = doc.add_paragraph('1. ЗАКАЗЧИК поручает, а ИСПОЛНИТЕЛЬ принимает на себя выполнение следующих работ:',
                          style='Normal')

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    for idx, work in enumerate(work_list, 1):
        p = doc.add_paragraph(f"1.{idx}. {work}.", style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Section 2
    p = doc.add_paragraph(f'2. Место проведения работ: {location}.', style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 3
    rubles, kopecks = divmod(total_cost, 1)
    rubles = int(rubles)
    kopecks = round(kopecks * 100)  # Округление до целых копеек
    sum_in_words = num2words(rubles, lang='ru', to='cardinal')

    p = doc.add_paragraph(
        f'3. Стоимость работ, согласно Приложению 1 к настоящему договору, составляет {total_cost:.2f} '
        f'({sum_in_words.capitalize()} белорусских рублей 00 коп.). Без НДС.',
        style='Normal'
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 4
    days_in_words = num2words(30, lang='ru', to='cardinal')  # Пример для 30 дней
    if completion[3] == "предоплата" and payment_term[3] != 100:
        p = doc.add_paragraph(
            f'4. Условия оплаты: Заказчик производит {payment_term[3]}% предоплату работ в течение 5(пяти) рабочих дней с момента подписания сторонами настоящего договора. {completion[4]}'
            f' 5(пяти) банковских дней после подписания сторонами Акта сдачи-приемки работ,'
            f' являющимся Приложением 2 к настоящему договору. Заказчик по своему усмотрению может произвести полную '
            f'предоплату работ.',
            style='Normal'
        )
    elif completion[3] == "предоплата" and payment_term[3] == 100:
        p = doc.add_paragraph(
            f'4. Условия оплаты: Заказчик производит {payment_term[3]}% предоплату работ в течение 5(пяти) рабочих дней с момента '
            f'подписания сторонами настоящего договора. ',
            style='Normal'
        )
    else:
        p = doc.add_paragraph(
            f'4. Условия оплаты: Заказчик производит оплату Работ в течение 5(пяти) банковских дней после подписания '
            f'сторонами Акта сдачи-приемки работ, являющегося '
            f'Приложением 2 к настоящему договору. Заказчик, по своему усмотрению, может произвести полную или '
            f'частичную предоплату работ.',
            style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 5
    if completion[3] == "предоплата":
        p = doc.add_paragraph(f'5. Окончание Работ: в течение 7(семи) рабочих дней с момента '
                              f'поступления на счет Исполнителя {payment_term[3]}% предоплаты.', style='Normal')
    else:
        p = doc.add_paragraph(f'5. Окончание Работ: в течение 7(семи) рабочих дней с момента подписания '
                              f'сторонами настоящего договора. Сроки выполнения Работ могут быть продлены по согласованию сторон.', style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 6
    p = doc.add_paragraph(
        '6. Заказчик в день получения акта сдачи-приемки работ обязан вернуть Исполнителю подписанный акт сдачи-приемки работ или мотивированный отказ от приемки работ.',
        style='Normal'
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Section 7
    p = doc.add_paragraph(
        '7. По всем остальным вопросам стороны руководствуются действующим законодательством Республики Беларусь.',
        style='Normal'
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Party details in a table
    recv=doc.add_paragraph('Реквизиты сторон:', style='Normal')
    recv.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    recv.runs[0].font.bold=True
    recv.paragraph_format.space_before = Pt(0)
    recv.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Pt(200)
    table.columns[1].width = Pt(200)

    # Fill in table
    row_cells = table.rows[0].cells
    # Set text and font size for the first cell (Исполнитель)
    p_ispolnitel = row_cells[0].paragraphs[0]

    # Создаем жирный Run для "ИП Панченко К.А."
    run_ispolnitel_bold = p_ispolnitel.add_run(f'{USER_COMPANY_SHORT}\n')
    run_ispolnitel_bold.font.size = Pt(9)
    run_ispolnitel_bold.bold = True

    # Добавляем остальной текст (обычным шрифтом)

    run_ispolnitel = p_ispolnitel.add_run(
        f"{OUR_REQUIZITS}"
    )
    run_ispolnitel.font.size = Pt(9)

    # Set text and font size for the second cell (Заказчик)
    p_zakazchik = row_cells[1].paragraphs[0]

    # Создаем жирный Run для названия заказчика
    run_zakazchik_bold = p_zakazchik.add_run(f'{customer[10]}\n')
    run_zakazchik_bold.font.size = Pt(9)
    run_zakazchik_bold.bold = True

    # Добавляем остальной текст (обычным шрифтом)
    run_zakazchik = p_zakazchik.add_run(
        f'{customer[5]}\n'
        f'IBAN: {customer[8]}\n'
        f'УНП: {customer[6]}, ОКПО: {customer[7]}\n'

    )
    run_zakazchik.font.size = Pt(9)

    # Set column widths
    table.columns[0].width = Pt(150)  # Example width for the first column
    table.columns[1].width = Pt(350)  # Example width for the second column


    # Signatures
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Set column widths
    table.columns[0].width = Pt(150)  # Example width for the first column
    table.columns[1].width = Pt(450)  # Example width for the second column

    cell1 = table.cell(0, 0)
    cell1.text = f'Исполнитель_________({USER_NAME})'
    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1.paragraphs[0].runs[0].font.size = Pt(9)

    cell2 = table.cell(0, 1)
    cell2.text = f'Заказчик_________({customer[4]})'
    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell2.paragraphs[0].runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    p = doc.add_paragraph()

    # Appendix 1
    p = doc.add_paragraph(f'Приложение 1 к ДОГОВОРУ № {contract_number} от {doc_date}', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph('Протокол согласования договорной цены', style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Appendix 1 Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    table.columns[0].width = Pt(50)
    table.columns[1].width = Pt(350)
    table.columns[2].width = Pt(100)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование работ'
    hdr_cells[2].text = 'Стоимость, руб.'

    for cell in hdr_cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Add data to the table
    num_works = len(work_list)
    price_per_work = total_cost / num_works  # Рассчитываем стоимость на каждую работу

    for idx, work in enumerate(work_list, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = work
        row_cells[2].text = f'{price_per_work:.2f}'  # Записываем рассчитанную стоимость
        row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Total row
    row_cells = table.add_row().cells
    row_cells[0].text = 'Итого:'
    cell = row_cells[1]
    cell.merge(row_cells[2])
    cell.text = f'{total_cost:.2f}'
    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)

    # Total sum
    # Total sum
    sum_total = total_cost
    paragraph = doc.add_paragraph(style='Normal')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выровнено по левому краю

    # Первая часть текста
    run = paragraph.add_run(f'Итого стоимость работ: ')
    run.font.bold = True
    run.font.size = Pt(10)

    run = paragraph.add_run(f'{sum_in_words.capitalize()} белорусских рублей 00 копеек.\n')
    run.font.size = Pt(10)

    # Вторая часть текста на новой строке
    run = paragraph.add_run('Без НДС.')
    run.font.size = Pt(10)

    # Signatures
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    cell1 = table.cell(0, 0)
    cell1.text = f'Исполнитель_________({USER_NAME})'
    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1.paragraphs[0].runs[0].font.size = Pt(9)

    cell2 = table.cell(0, 1)
    cell2.text = f'Заказчик___________({customer[4]})'
    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell2.paragraphs[0].runs[0].font.size = Pt(9)


    # Save document with contract number

    file_path = filedialog.asksaveasfilename(
        defaultextension=f".docx",
        filetypes=[("Word files", "*.docx")],
        initialfile=f'contract_{contract_number}.docx'
    )

    #file_name = f'../docs_out/contract_{contract_number}.docx'
    if file_path:
        # Сохраняем документ
        doc.save(file_path)

        # Получаем папку, в которую пользователь сохранил файл
        saved_folder = os.path.dirname(file_path)
        print("Папка, выбранная пользователем:", saved_folder)

        return file_path  # или return saved_folder, если тебе именно папка нужна
    else:
        return None

